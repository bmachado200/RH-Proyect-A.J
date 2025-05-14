#   File "C:\Users\alex2\OneDrive\Escritorio\RH-NEWCONCEPT\LoaderContrato.py", line 200, in <module>
 #   procesar_documentos()
  #File "C:\Users\alex2\OneDrive\Escritorio\RH-NEWCONCEPT\LoaderContrato.py", line 137, in procesar_documentos
   # files = [f for f in os.listdir(DOCS_DIR) if f.lower().endswith('.docx')]
    #                    ^^^^^^^^^^^^^^^^^^^^
#FileNotFoundError: [WinError 3] The system cannot find the path specified: './HR/contrato_colectivo'

import os
from docx import Document
import chromadb
from typing import List
from sentence_transformers import SentenceTransformer
import torch
from tqdm import tqdm
import logging
from typing import List  # Add this with other imports
from datetime import datetime
import re

# Configuración específica
DOCS_DIR = "./HR/documentos_internos"
DB_DIR = "./chroma_db_internos"
COLLECTION_NAME = "documentos_internos"
CHUNK_SIZE = 2000
CHUNK_OVERLAP = 300
MIN_CHUNK_LENGTH = 50

# Configuración de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('loader_internos.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Configuración GPU
device = "cuda" if torch.cuda.is_available() else "cpu"
if device == "cuda":
    torch.backends.cudnn.benchmark = True
    logger.info(f"GPU activada: {torch.cuda.get_device_name(0)}")

# Cargar modelo de embeddings
embedding_model = SentenceTransformer(
    'BAAI/bge-m3',
    device=device,
    cache_folder='./model_cache'
)
embedding_model.max_seq_length = 8192

def extraer_texto(file_path: str) -> str:
    """Extrae texto estructurado de documentos internos"""
    try:
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                text_parts.append(f"\n\nSECCIÓN: {para.text.upper()}\n")
            else:
                text_parts.append(para.text)
        
        for table in doc.tables:
            text_parts.append("\n[TABLA]")
            for row in table.rows:
                text_parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            text_parts.append("[/TABLA]\n")
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extrayendo {file_path}: {e}")
        raise

def crear_chunks(texto: str) -> List[str]:
    """Divide el texto preservando secciones importantes"""
    chunks = []
    
    # Dividir por secciones primero
    sections = re.split(r'\n\nSECCIÓN: (POLÍTICA|PROCEDIMIENTO|BENEFICIO|NORMATIVA)', texto)
    
    for i in range(1, len(sections), 2):
        section_title = sections[i]
        section_content = sections[i+1]
        full_section = f"SECCIÓN: {section_title}\n{section_content}"
        
        if len(full_section) <= CHUNK_SIZE * 1.5:
            chunks.append(full_section)
            continue
            
        # Dividir secciones grandes
        lines = [line.strip() for line in full_section.split('\n') if line.strip()]
        current_chunk = []
        current_length = 0
        
        for line in lines:
            line_length = len(line)
            if current_length + line_length > CHUNK_SIZE and current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = current_chunk[-CHUNK_OVERLAP:]
                current_length = sum(len(l) for l in current_chunk)
            current_chunk.append(line)
            current_length += line_length
        
        if current_chunk:
            chunks.append("\n".join(current_chunk))
    
    return [chunk for chunk in chunks if len(chunk) >= MIN_CHUNK_LENGTH]

def generar_embeddings(textos: List[str]) -> List[List[float]]:
    """Genera embeddings para los textos"""
    try:
        with torch.no_grad(), torch.amp.autocast(device_type=device):
            instruction = "Representar documento interno para recuperación: "
            textos = [instruction + text for text in textos]
            return embedding_model.encode(
                textos,
                batch_size=128 if device == "cuda" else 16,
                show_progress_bar=False,
                normalize_embeddings=True
            ).tolist()
    except Exception as e:
        logger.error(f"Error generando embeddings: {e}")
        raise

def procesar_documentos():
    """Procesa todos los documentos en DOCS_DIR"""
    try:
        client = chromadb.PersistentClient(path=DB_DIR)
        collection = client.get_or_create_collection(
            name=COLLECTION_NAME,
            metadata={
                "hnsw:space": "cosine",
                "tipo": "documentos_internos"
            }
        )
        
        existing_ids = set(collection.get(include=[], limit=100000)['ids'])
        files = [f for f in os.listdir(DOCS_DIR) if f.lower().endswith('.docx')]
        
        if not files:
            logger.warning("No se encontraron documentos .docx")
            return
        
        logger.info(f"Procesando {len(files)} documentos internos")
        
        total_chunks = 0
        for filename in tqdm(files, desc="Procesando"):
            try:
                filepath = os.path.join(DOCS_DIR, filename)
                text = extraer_texto(filepath)
                chunks = crear_chunks(text)
                
                new_chunks = []
                chunk_ids = []
                for i, chunk in enumerate(chunks):
                    chunk_id = f"{filename}_chunk_{i}"
                    if chunk_id not in existing_ids:
                        new_chunks.append(chunk)
                        chunk_ids.append(chunk_id)
                
                if not new_chunks:
                    continue
                
                embeddings = generar_embeddings(new_chunks)
                
                metadatos = [{
                    "fuente": filename,
                    "tipo": "interno",
                    "seccion": chunk.split("SECCIÓN:")[1].split("\n")[0] if "SECCIÓN:" in chunk else "general",
                    "longitud": len(chunk),
                    "fecha_procesamiento": datetime.now().isoformat()
                } for chunk in new_chunks]
                
                collection.add(
                    documents=new_chunks,
                    embeddings=embeddings,
                    metadatas=metadatos,
                    ids=chunk_ids
                )
                total_chunks += len(new_chunks)
                
            except Exception as e:
                logger.error(f"Error procesando {filename}: {e}")
                continue
        
        logger.info(f"Proceso completado. Chunks añadidos: {total_chunks}")
        logger.info(f"Total en colección: {collection.count()}")
        
    except Exception as e:
        logger.error(f"Error fatal: {e}")
        raise

if __name__ == "__main__":
    logger.info("Iniciando carga de documentos internos")
    start_time = datetime.now()
    
    try:
        # Calentamiento inicial
        embedding_model.encode(["Documento de prueba"] * 10)
        
        procesar_documentos()
    finally:
        duration = datetime.now() - start_time
        logger.info(f"Tiempo total: {duration}")
        
        if device == "cuda":
            torch.cuda.empty_cache()